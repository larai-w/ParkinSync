import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_april_word_log():
    doc = Document()
    
    # ページを横向き（Landscape）に設定
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    new_width, new_height = section.page_height, section.page_width
    section.page_width = new_width
    section.page_height = new_height
    
    # タイトル追加
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("ParkinSync Caregiver Log - April 2026")
    run.font.size = Pt(18)
    run.font.bold = True
    
    # 10列×31行（ヘッダー1行 + データ30行）の表を作成
    headers = ["Date", "Day", "Morning Meds", "Lunch Meds", "Evening Meds", "20:30 / Bedtime", "Bowel / Movicol", "Falls (Count)", "Emergency Call", "Daily Notes"]
    table = doc.add_table(rows=1, cols=10)
    table.style = 'Table Grid'
    
    # ヘッダーの文字入れ
    hdr_cells = table.rows[0].cells
    for i, header_text in enumerate(headers):
        hdr_cells[i].text = header_text
        hdr_cells[i].paragraphs[0].runs[0].font.bold = True
        hdr_cells[i].paragraphs[0].runs[0].font.size = Pt(9)
    
    days_of_week = ["Wed", "Thu", "Fri", "Sat", "Sun", "Mon", "Tue"]
    
    # 30日分の行を追加してデータを埋める
    for day in range(1, 31):
        wday = days_of_week[(day - 1) % 7]
        is_highlight = wday in ["Wed", "Sat"]
        
        suffix = "st" if day == 1 else "nd" if day == 2 else "rd" if day == 3 else "th"
        date_str = f"{day}{suffix}"
        
        row_cells = table.add_row().cells
        row_cells[0].text = date_str
        row_cells[1].text = wday
        
        # フォントサイズ調整とハイライト設定
        for i in range(10):
            p = row_cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if len(p.runs) > 0:
                p.runs[0].font.size = Pt(8)
                if is_highlight and i < 2:
                    p.runs[0].font.bold = True
                    p.runs[0].font.color.rgb = RGBColor(192, 57, 43) # 赤文字
    
    # フッターの記録基準を追加
    doc.add_paragraph("\n[Recording Standards]")
    
    p1 = doc.add_paragraph("• Condition (1-5): 1 = Severe OFF / 5 = Fully ON")
    if p1.runs:
        p1.runs[0].font.size = Pt(9)
        
    p2 = doc.add_paragraph("• Post-Medication Longitudinal Tracking (Condition 1-5): 30m [ ] / 1h [ ] / 2h [ ] / 3h [ ] / 4h [ ] / 5h [ ]")
    if p2.runs:
        p2.runs[0].font.size = Pt(9)
    
    # ファイル保存
    output_path = os.path.join("design", "log_template_2026_04.docx")
    doc.save(output_path)
    print(f"[SUCCESS] Word document generated at: {output_path}")

if __name__ == "__main__":
    create_april_word_log()